# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY = RGBColor(0x1F, 0x3A, 0x5F)
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GREY = RGBColor(0x44, 0x44, 0x44)

doc = Document()

# ---- base style ----
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.get_or_add_rFonts()
rfonts.set(qn('w:cs'), 'Arial')
style.paragraph_format.space_after = Pt(6)

# margins
sec = doc.sections[0]
sec.top_margin = Cm(2)
sec.bottom_margin = Cm(2)
sec.left_margin = Cm(2.2)
sec.right_margin = Cm(2.2)

def set_rtl(paragraph):
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)

def run_rtl(run):
    rPr = run._r.get_or_add_rPr()
    rtl = OxmlElement('w:rtl')
    rPr.append(rtl)

def add_par(text, size=12, bold=False, color=None, align=WD_ALIGN_PARAGRAPH.RIGHT,
            space_before=0, space_after=6, italic=False):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    run_rtl(r)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    if color is not None:
        r.font.color.rgb = color
    return p

def add_heading(text):
    p = add_par(text, size=13.5, bold=True, color=NAVY, space_before=10, space_after=4)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '2')
    bottom.set(qn('w:color'), 'B8860B')
    pbdr.append(bottom)
    pPr.append(pbdr)
    return p

def add_bullet(text):
    p = doc.add_paragraph()
    set_rtl(p)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("• " + text)
    run_rtl(r)
    r.font.size = Pt(12)
    r.font.name = 'Arial'
    r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
    return p

# ================= TITLE =================
add_par("عقد شراكة والتزام", size=22, bold=True, color=NAVY,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_before=6, space_after=2)
add_par("ــــــــــــــــــــــــــــ", size=14, color=GOLD,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

# ================= PARTIES =================
add_par("بين الطرفين الموقّعين أدناه:", bold=True, space_after=6)

add_par("الشركة: Depart Travel Services، شركة مؤسّسة قانونيًا وفقًا للتشريع "
        "التونسي، ويمثّلها ممثّلها القانوني السيّد العمراني أيمن، صاحب بطاقة "
        "التعريف الوطنية عدد 09107429 الصادرة بتاريخ 29/09/2022، ويُشار إليها "
        "فيما يلي بـ «الوكالة».")

add_par("والسيّد الفاتح الهليّل، صاحب بطاقة التعريف الوطنية عدد 08212752 "
        "الصادرة بتاريخ 23/12/2009، القاطن بـ: بني خداش – مدنين، ويُشار إليه "
        "فيما يلي بـ «المستثمر».")

add_par("ويُشار إليهما معًا بـ «الطرفين».", space_after=8)

# ================= ARTICLES =================
add_heading("الفصل 1 – موضوع العقد")
add_par("يهدف هذا العقد إلى تحديد حقوق والتزامات الطرفين فيما يتعلق بالسيّارة التالية:")
add_par("رقم التسجيل: 262 تونس 1043", bold=True)
add_par("المُقتناة عن طريق الإيجار المالي (Leasing) باسم Depart Travel Services.")
add_par("يقرّ الطرفان بأنّ السيّارة مسجّلة قانونيًا باسم الوكالة طيلة مدّة الإيجار "
        "المالي، غير أنّها مُقتناة لفائدة المصلحة الاقتصادية للسيّد الفاتح الهليّل "
        "وفق شروط هذا العقد.")

add_heading("الفصل 2 – مساهمة المستثمر")
add_par("يصرّح السيّد الفاتح الهليّل بأنّه دفع للوكالة مبلغ سبعة آلاف دينار تونسي "
        "(7000 د.ت) على سبيل التسبقة لاقتناء السيّارة. وتقرّ الوكالة بتسلّم هذا المبلغ.")

add_heading("الفصل 3 – استغلال السيّارة")
add_par("تُستغلّ السيّارة حصريًا من قبل الوكالة في إطار نشاطها المهني. وتتولّى الوكالة:")
add_bullet("إدارة الاستغلال؛")
add_bullet("البحث عن الحرفاء؛")
add_bullet("تنظيم المهمّات؛")
add_bullet("المتابعة الإدارية لاستعمال السيّارة.")

add_heading("الفصل 4 – خلاص أقساط الإيجار المالي")
add_par("تلتزم الوكالة طيلة مدّة الإيجار المالي بخلاص جميع الأقساط المستحقّة لشركة "
        "الإيجار المالي على نفقتها الخاصّة وفي الآجال المحدّدة. ولا يمكن بأيّ حال من "
        "الأحوال تحميل عدم خلاص الأقساط للسيّد الفاتح الهليّل.")

add_heading("الفصل 5 – المصاريف المتعلّقة بالسيّارة")
add_par("تُوزّع المصاريف على النحو التالي:")
add_par("على عاتق السيّد الفاتح الهليّل:", bold=True, color=NAVY, space_before=2)
add_bullet("تأمين السيّارة؛")
add_bullet("الوقود؛")
add_bullet("الصيانة العادية؛")
add_bullet("الإصلاحات الميكانيكية؛")
add_bullet("العجلات المطّاطية؛")
add_bullet("الفحوصات الفنية؛")
add_bullet("الأداءات وكلّ المصاريف الأخرى المتعلّقة بملكية السيّارة وصيانتها.")
add_par("على عاتق الوكالة:", bold=True, color=NAVY, space_before=2)
add_bullet("أقساط الإيجار المالي؛")
add_bullet("المصاريف المرتبطة مباشرة بالاستغلال التجاري للسيّارة في إطار نشاطها "
           "(السائق، مصاريف الاستغلال، المعاليم، ركن السيّارة، إلخ، حسب حاجيات النشاط).")

add_heading("الفصل 6 – تحديد الأرباح")
add_par("يتمثّل الربح السنوي في المداخيل المتأتّية من استغلال السيّارة، منقوصًا منها "
        "مجموع مصاريف الاستغلال، بما في ذلك خاصّة:")
add_bullet("أقساط الإيجار المالي؛")
add_bullet("أجور أو مكافآت السائقين؛")
add_bullet("المصاريف الإدارية المرتبطة مباشرة بالسيّارة؛")
add_bullet("مصاريف الاستغلال.")
add_par("أمّا المصاريف التي يتحمّلها شخصيًا السيّد الفاتح الهليّل (التأمين، الوقود، "
        "الصيانة، إلخ) فلا تتحمّلها الوكالة. ويُعدّ كشف تفصيلي بالمداخيل والمصاريف كلّ سنة.")

add_heading("الفصل 7 – صرف الأرباح")
add_par("بعد إعداد الميزانية السنوية للسيّارة، تصرف الوكالة للسيّد الفاتح الهليّل "
        "كامل الربح الصافي المتأتّي من هذه السيّارة. ويتمّ الخلاص في أجل أقصاه ثلاثون "
        "(30) يومًا من تاريخ ختم السنة المالية المعنية. ويمكن للسيّد الفاتح الهليّل "
        "الاطّلاع على المؤيّدات المحاسبية المتعلّقة بالسيّارة.")

add_heading("الفصل 8 – غياب الأرباح")
add_par("إذا لم تحقّق السيّارة أيّ ربح خلال سنة ما أو سجّلت خسارة:")
add_bullet("تواصل الوكالة خلاص أقساط الإيجار المالي؛")
add_bullet("تواصل الوكالة استعمال السيّارة في نشاطها؛")
add_bullet("لا يمكن للسيّد الفاتح الهليّل المطالبة بأيّ ربح عن هذه الفترة؛")
add_bullet("لا يكون السيّد الفاتح الهليّل ملزمًا بإرجاع أقساط الإيجار المالي التي خلّصتها الوكالة.")

add_heading("الفصل 9 – نقل الملكية")
add_par("بعد الخلاص الكامل لجميع أقساط الإيجار المالي، ومع مراعاة الإجراءات القانونية "
        "وكذلك الموافقة المحتملة لشركة الإيجار المالي، تلتزم الوكالة بنقل ملكية السيّارة "
        "باسم السيّد الفاتح الهليّل دون المطالبة بأيّ مقابل إضافي، باستثناء المصاريف "
        "الإدارية والأداءات المتعلّقة بتغيير المالك، ما لم يتّفق الطرفان على خلاف ذلك.")

add_heading("الفصل 10 – التأمين والمسؤولية")
add_par("طيلة مدّة العقد:")
add_bullet("يجب أن تبقى السيّارة مؤمَّنة طبقًا للتشريع التونسي؛")
add_bullet("يتحمّل كلّ طرف الالتزامات المنوطة بعهدته بمقتضى هذا العقد؛")
add_bullet("في صورة وقوع حادث، يتعاون الطرفان للحفاظ على حقوقهما.")

add_heading("الفصل 11 – المدّة")
add_par("يدخل هذا العقد حيّز التنفيذ من تاريخ إمضائه، ويبقى ساري المفعول إلى حين "
        "النقل النهائي للسيّارة باسم السيّد الفاتح الهليّل، ما لم يتمّ فسخه بالتراضي "
        "أو في حالة القوّة القاهرة.")

add_heading("الفصل 12 – فضّ النزاعات")
add_par("يلتزم الطرفان بالبحث عن حلّ ودّي لكلّ خلاف. وفي حال تعذّر الاتّفاق الودّي، "
        "تكون المحاكم المختصّة بالجمهورية التونسية وحدها المختصّة.")

add_heading("الفصل 13 – أحكام ختامية")
add_par("كلّ تعديل لهذا العقد يجب أن يكون موضوع ملحق كتابي مُمضى من الطرفين. وحُرّر "
        "هذا العقد في نسختين أصليتين، نسخة لكلّ طرف.")

# ================= SIGNATURES =================
add_par("حُرّر بجربة في 15 / 07 / 2026", bold=True, space_before=14, space_after=14,
        align=WD_ALIGN_PARAGRAPH.CENTER)

# signature table
table = doc.add_table(rows=1, cols=2)
table.alignment = 1
table.autofit = True
cellL, cellR = table.rows[0].cells

def fill_cell(cell, lines):
    cell.paragraphs[0].clear() if False else None
    first = True
    for text, bold in lines:
        if first:
            p = cell.paragraphs[0]
            first = False
        else:
            p = cell.add_paragraph()
        set_rtl(p)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(8)
        r = p.add_run(text)
        run_rtl(r)
        r.bold = bold
        r.font.size = Pt(12)
        r.font.name = 'Arial'
        r._element.rPr.rFonts.set(qn('w:cs'), 'Arial')
        if bold:
            r.font.color.rgb = NAVY

# Right cell = الوكالة (RTL, right first)
fill_cell(cellR, [
    ("الوكالة", True),
    ("Depart Travel Services", False),
    ("الاسم: العمراني أيمن", False),
    ("الإمضاء:", False),
    ("الختم:", False),
])
fill_cell(cellL, [
    ("المستثمر", True),
    ("السيّد الفاتح الهليّل", False),
    ("بطاقة التعريف الوطنية عدد: 08212752", False),
    ("الإمضاء:", False),
    ("", False),
])

# borders on table
tbl = table._tbl
tblPr = tbl.tblPr
borders = OxmlElement('w:tblBorders')
for edge in ('top','left','bottom','right','insideH','insideV'):
    e = OxmlElement(f'w:{edge}')
    e.set(qn('w:val'), 'single')
    e.set(qn('w:sz'), '4')
    e.set(qn('w:color'), 'CCCCCC')
    borders.append(e)
tblPr.append(borders)

doc.save('/home/user/qzdsq/Contrat_Partenariat_AR.docx')
print("saved")
